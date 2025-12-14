"""
Serwis generowania i eksportu raportów
"""
import sys
import os
import base64
import json
from typing import List, Dict, Optional
from datetime import datetime
import markdown
from io import BytesIO

# Dodaj ścieżkę do projektu
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from models.classification_result import ClassificationResult
from models.category_key import CategoryKey
from services.visualization_service import VisualizationService
from services.gemini_service import GeminiService

# Importy dla eksportu
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ReportService:
    """Serwis generowania raportów analitycznych"""
    
    def __init__(self):
        self.reports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'reports')
        os.makedirs(self.reports_dir, exist_ok=True)
        self.visualization_service = VisualizationService()
        self.gemini_service = GeminiService()
    
    def generate_report(
        self,
        classification_results: List[ClassificationResult],
        category_key: CategoryKey,
        brand_name: str,
        job_id: str,
        start_date: str,
        end_date: str
    ) -> Dict[str, str]:
        """
        Główna metoda generowania raportu
        Zwraca dict z ścieżkami do plików: {"html": path, "markdown": content}
        """
        # 1. Generuj wykresy
        chart_paths = self.visualization_service.generate_all_charts(classification_results, job_id)
        
        # 2. Przygotuj statystyki
        stats = self._calculate_statistics(classification_results, category_key)
        
        # 3. Generuj treść przez Gemini
        report_markdown = self._generate_report_content_with_gemini(
            classification_results, category_key, brand_name, stats, start_date, end_date
        )
        
        # 4. Finalizuj raport - wstaw wykresy
        final_markdown = self._embed_charts_in_markdown(report_markdown, chart_paths, job_id)
        
        # 5. Konwertuj do HTML
        html_content = self._convert_markdown_to_html(final_markdown, chart_paths, job_id)
        
        # 6. Zapisz HTML
        html_path = os.path.join(self.reports_dir, f'report_{job_id}.html')
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # 7. Zapisz Markdown (dla eksportu DOCX)
        markdown_path = os.path.join(self.reports_dir, f'report_{job_id}.md')
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(final_markdown)
        
        return {
            "html": html_path,
            "markdown": markdown_path,
            "markdown_content": final_markdown,
            "charts": chart_paths
        }
    
    def export_to_docx(self, markdown_content: str, chart_paths: Dict[str, str], job_id: str) -> Optional[str]:
        """Eksportuje raport Markdown do DOCX"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx nie jest zainstalowany. Zainstaluj: pip install python-docx")
        
        docx_path = os.path.join(self.reports_dir, f'report_{job_id}.docx')
        
        # Utwórz dokument DOCX
        doc = Document()
        
        def process_text_with_formatting(text: str, para):
            """Przetwarza tekst z markdown formatting i dodaje do paragrafu"""
            if not text:
                return
            
            # Obsługa **bold** - dzielimy tekst na części przed i po **
            parts = text.split('**')
            for i, part in enumerate(parts):
                if part:  # Pomiń puste części
                    if i % 2 == 0:
                        # Parzyste indeksy = zwykły tekst
                        para.add_run(part)
                    else:
                        # Nieparzyste indeksy = tekst do pogrubienia
                        para.add_run(part).bold = True
        
        # Parsuj Markdown linia po linii
        lines = markdown_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            original_line = line
            line_stripped = line.strip()
            
            if not line_stripped:
                i += 1
                continue
            
            # Nagłówki - usuwamy formatowanie (DOCX add_heading nie obsługuje formatowania w tekście)
            if line_stripped.startswith('# '):
                heading_text = line_stripped[2:].strip()
                heading_text = heading_text.replace('**', '')  # Usuń markdown bold
                doc.add_heading(heading_text, level=1)
            elif line_stripped.startswith('## '):
                heading_text = line_stripped[3:].strip()
                heading_text = heading_text.replace('**', '')
                doc.add_heading(heading_text, level=2)
            elif line_stripped.startswith('### '):
                heading_text = line_stripped[4:].strip()
                heading_text = heading_text.replace('**', '')
                doc.add_heading(heading_text, level=3)
            # Listy
            elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                para = doc.add_paragraph(style='List Bullet')
                list_text = line_stripped[2:].strip()
                process_text_with_formatting(list_text, para)
            elif line_stripped[0].isdigit() and '. ' in line_stripped[:5]:
                # Lista numerowana
                para = doc.add_paragraph(style='List Number')
                list_text = line_stripped.split('. ', 1)[1] if '. ' in line_stripped else line_stripped
                process_text_with_formatting(list_text.strip(), para)
            # Obrazy
            elif line_stripped.startswith('![') and '](' in line_stripped:
                # Format: ![alt](path)
                alt_end = line_stripped.find('](')
                path_start = alt_end + 2
                path_end = line_stripped.find(')', path_start)
                if path_end > path_start:
                    img_path = line_stripped[path_start:path_end]
                    # Sprawdź czy to ścieżka do wykresu
                    if img_path.startswith('/static/charts/'):
                        chart_filename = os.path.basename(img_path)
                        # Znajdź odpowiedni wykres w chart_paths
                        for chart_type, chart_path in chart_paths.items():
                            if chart_filename.endswith(os.path.basename(chart_path)):
                                if os.path.exists(chart_path):
                                    doc.add_picture(chart_path, width=Inches(6))
                                    break
            # Zwykły tekst z formatowaniem
            else:
                para = doc.add_paragraph()
                process_text_with_formatting(line_stripped, para)
            
            i += 1
        
        # Dodaj wykresy na końcu jeśli nie zostały dodane w treści
        doc.add_page_break()
        doc.add_heading('Wykresy', level=2)
        for chart_type, chart_path in chart_paths.items():
            if chart_path and os.path.exists(chart_path):
                chart_names = {
                    'bar': 'Rozkład kategorii i sentymentu',
                    'sentiment_pie': 'Rozkład sentymentu',
                    'categories_pie': 'Rozkład kategorii'
                }
                doc.add_heading(chart_names.get(chart_type, chart_type), level=3)
                doc.add_picture(chart_path, width=Inches(6))
        
        doc.save(docx_path)
        return docx_path
    
    def _calculate_statistics(self, classification_results: List[ClassificationResult], category_key: CategoryKey) -> Dict:
        """Oblicza statystyki z wyników klasyfikacji"""
        total = len(classification_results)
        
        # Sentiment
        sentiment_counts = {'pozytywny': 0, 'neutralny': 0, 'negatywny': 0}
        for result in classification_results:
            sentiment_counts[result.sentiment] += 1
        
        # Kategorie
        category_counts = {}
        for result in classification_results:
            category_counts[result.category] = category_counts.get(result.category, 0) + 1
        
        # Top kategorie
        top_categories = sorted(category_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        
        return {
            'total_comments': total,
            'sentiment_counts': sentiment_counts,
            'sentiment_percentages': {
                k: round(v / total * 100, 1) if total > 0 else 0 
                for k, v in sentiment_counts.items()
            },
            'category_counts': category_counts,
            'top_categories': top_categories,
            'categories_count': len(category_counts)
        }
    
    def _generate_report_content_with_gemini(
        self,
        classification_results: List[ClassificationResult],
        category_key: CategoryKey,
        brand_name: str,
        stats: Dict,
        start_date: str,
        end_date: str
    ) -> str:
        """Generuje treść raportu przez Gemini"""
        
        # Przygotuj dane dla Gemini
        categories_summary = []
        for cat in category_key.categories:
            cat_name = cat.get('aspekt', '')
            count = stats['category_counts'].get(cat_name, 0)
            categories_summary.append(f"- {cat_name}: {count} komentarzy")
        
        # Przykładowe komentarze (top 3 z każdej kategorii)
        examples_by_category = {}
        for result in classification_results:
            if result.category not in examples_by_category:
                examples_by_category[result.category] = []
            if len(examples_by_category[result.category]) < 3:
                examples_by_category[result.category].append(result.comment_text[:200])
        
        examples_text = "\n".join([
            f"**{cat}**:\n" + "\n".join([f"- {ex}" for ex in examples[:2]])
            for cat, examples in list(examples_by_category.items())[:5]
        ])
        
        prompt = f"""Jesteś analitykiem Customer Experience przygotowującym profesjonalny raport analityczny dla marki "{brand_name}".

**Dane wejściowe:**
- Okres analizy: {start_date} - {end_date}
- Łączna liczba przeanalizowanych komentarzy: {stats['total_comments']}
- Liczba kategorii (aspektów): {stats['categories_count']}

**Rozkład sentymentu:**
- Pozytywny: {stats['sentiment_percentages']['pozytywny']}% ({stats['sentiment_counts']['pozytywny']} komentarzy)
- Neutralny: {stats['sentiment_percentages']['neutralny']}% ({stats['sentiment_counts']['neutralny']} komentarzy)
- Negatywny: {stats['sentiment_percentages']['negatywny']}% ({stats['sentiment_counts']['negatywny']} komentarzy)

**Kategorie i liczba komentarzy:**
{chr(10).join(categories_summary)}

**Przykładowe komentarze:**
{examples_text}

**Najczęściej występujące kategorie:**
{chr(10).join([f"- {cat}: {count} komentarzy" for cat, count in stats['top_categories']])}

**Zadanie:**
Wygeneruj profesjonalny raport analityczny w formacie Markdown zawierający:

1. **Podsumowanie wykonawcze** (Executive Summary)
   - Krótkie podsumowanie głównych wniosków (2-3 akapity)

2. **Analiza statystyczna**
   - Interpretacja rozkładu sentymentu
   - Analiza najważniejszych kategorii
   - Trendy i wzorce

3. **Analiza kategorii**
   - Dla każdej z top 5 kategorii: krótka analiza, co mówią komentarze, główne tematy

4. **Wnioski i rekomendacje**
   - Główne wnioski z analizy
   - Konkretne rekomendacje dla marki (3-5 punktów)

**Wymagania:**
- Użyj profesjonalnego języka biznesowego
- Bądź konkretny i opieraj się na danych
- Wykorzystaj placeholdery [WYKRES_SENTIMENT] i [WYKRES_KATEGORIE] tam gdzie odpowiednie
- Formatuj w Markdown (nagłówki ##, listy, pogrubienia)
- Raport powinien mieć około 800-1200 słów

Zacznij od tytułu: # Raport Analizy Komentarzy - {brand_name}"""

        # Wywołaj Gemini
        response = self.gemini_service.flash_model.generate_content(prompt)
        report_text = response.text.strip()
        
        return report_text
    
    def _embed_charts_in_markdown(self, markdown_content: str, chart_paths: Dict[str, str], job_id: str) -> str:
        """Wstawia wykresy do Markdown"""
        content = markdown_content
        
        # Zamień placeholdery na obrazy
        if chart_paths.get('sentiment_pie') and os.path.exists(chart_paths['sentiment_pie']):
            chart_url = f"/static/charts/sentiment_pie_{job_id}.png"
            content = content.replace('[WYKRES_SENTIMENT]', f'![Rozkład sentymentu]({chart_url})')
        
        if chart_paths.get('categories_pie') and os.path.exists(chart_paths['categories_pie']):
            chart_url = f"/static/charts/categories_pie_{job_id}.png"
            content = content.replace('[WYKRES_KATEGORIE]', f'![Rozkład kategorii]({chart_url})')
        
        if chart_paths.get('bar') and os.path.exists(chart_paths['bar']):
            chart_url = f"/static/charts/bar_chart_{job_id}.png"
            # Jeśli nie ma placeholder, dodaj na końcu sekcji statystycznej
            if '[WYKRES_BAR]' in content:
                content = content.replace('[WYKRES_BAR]', f'![Rozkład kategorii i sentymentu]({chart_url})')
            elif '## Analiza statystyczna' in content:
                content = content.replace(
                    '## Analiza statystyczna',
                    f'## Analiza statystyczna\n\n![Rozkład kategorii i sentymentu]({chart_url})\n'
                )
        
        return content
    
    def _convert_markdown_to_html(self, markdown_content: str, chart_paths: Dict[str, str], job_id: str) -> str:
        """Konwertuje Markdown do HTML z osadzonymi wykresami"""
        # Konwertuj Markdown do HTML
        html_body = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
        
        # Konwertuj ścieżki wykresów na base64 dla PDF
        for chart_type, chart_path in chart_paths.items():
            if chart_path and os.path.exists(chart_path):
                # Wczytaj obraz i konwertuj na base64
                with open(chart_path, 'rb') as img_file:
                    img_data = base64.b64encode(img_file.read()).decode('utf-8')
                    img_ext = os.path.splitext(chart_path)[1][1:]  # png, jpg, etc.
                    img_base64 = f"data:image/{img_ext};base64,{img_data}"
                    
                    # Zamień ścieżki na base64 w HTML
                    chart_filename = os.path.basename(chart_path)
                    html_body = html_body.replace(
                        f'src="/static/charts/{chart_filename}"',
                        f'src="{img_base64}"'
                    )
                    html_body = html_body.replace(
                        f'![',
                        f'<img src="{img_base64}" alt="'
                    )
        
        # Pełny HTML z CSS
        html_template = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Raport Analizy - {job_id}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 5px;
        }}
        h3 {{
            color: #7f8c8d;
            margin-top: 20px;
        }}
        img {{
            max-width: 100%;
            height: auto;
            margin: 20px 0;
            border: 1px solid #ddd;
            border-radius: 4px;
        }}
        ul, ol {{
            margin: 10px 0;
            padding-left: 30px;
        }}
        strong {{
            color: #2c3e50;
        }}
        .meta {{
            background: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""
        
        return html_template
    
    def _convert_chart_to_base64(self, chart_path: str) -> str:
        """Konwertuje wykres PNG do base64"""
        if not os.path.exists(chart_path):
            return ""
        
        with open(chart_path, 'rb') as img_file:
            img_data = base64.b64encode(img_file.read()).decode('utf-8')
            img_ext = os.path.splitext(chart_path)[1][1:]
            return f"data:image/{img_ext};base64,{img_data}"
